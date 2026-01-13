"""
文档处理模块

负责文档解析、文本提取和智能分块
支持PDF、TXT、DOCX、Markdown格式
"""

import re
import uuid
import logging
from pathlib import Path
from dataclasses import dataclass, field
from datetime import datetime
from typing import List, Dict, Optional, Tuple

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """文档数据类"""
    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    filename: str = ""
    content: str = ""
    format: str = ""
    metadata: Dict = field(default_factory=dict)
    created_at: datetime = field(default_factory=datetime.now)
    chunk_count: int = 0


@dataclass
class Chunk:
    """文本块数据类"""
    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    document_id: str = ""
    content: str = ""
    start_pos: int = 0
    end_pos: int = 0
    metadata: Dict = field(default_factory=dict)


class DocumentParseError(Exception):
    """文档解析错误"""
    pass


class DocumentLoader:
    """
    文档加载器
    
    负责加载和解析各种格式的文档
    """
    
    def __init__(self, supported_formats: List[str] = None):
        self.supported_formats = supported_formats or [".pdf", ".txt", ".docx", ".md"]
    
    def load(self, file_path: str) -> Document:
        """加载文档"""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        suffix = path.suffix.lower()
        if suffix not in self.supported_formats:
            raise DocumentParseError(f"不支持的文件格式: {suffix}")
        
        parsers = {
            ".pdf": self._parse_pdf,
            ".txt": self._parse_txt,
            ".docx": self._parse_docx,
            ".md": self._parse_markdown,
        }
        
        try:
            content = parsers[suffix](str(path))
        except Exception as e:
            raise DocumentParseError(f"解析文档失败: {e}")
        
        return Document(
            filename=path.name,
            content=content,
            format=suffix,
            metadata={
                "file_path": str(path),
                "file_size": path.stat().st_size,
            }
        )
    
    def _parse_pdf(self, file_path: str) -> str:
        """解析PDF文档"""
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n\n".join(text_parts)
        except ImportError:
            logger.warning("pypdf未安装，尝试使用备选方案")
            return self._parse_txt(file_path)
    
    def _parse_txt(self, file_path: str) -> str:
        """解析纯文本文档"""
        encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise DocumentParseError("无法识别文件编码")
    
    def _parse_docx(self, file_path: str) -> str:
        """解析Word文档"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            return "\n\n".join(paragraphs)
        except ImportError:
            logger.warning("python-docx未安装")
            raise DocumentParseError("需要安装python-docx来解析DOCX文件")
    
    def _parse_markdown(self, file_path: str) -> str:
        """解析Markdown文档"""
        content = self._parse_txt(file_path)
        content = re.sub(r"```[\s\S]*?```", "", content)
        content = re.sub(r"`[^`]+`", "", content)
        content = re.sub(r"!\[.*?\]\(.*?\)", "", content)
        content = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", content)
        content = re.sub(r"^#+\s+", "", content, flags=re.MULTILINE)
        return content


class TextChunker:
    """
    文本分块器基类
    """
    
    def chunk(self, text: str, doc_id: str) -> List[Chunk]:
        raise NotImplementedError


class SemanticChunker(TextChunker):
    """
    语义分块器 - 创新点
    
    基于语义边界进行智能分块，避免在句子中间切分
    """
    
    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 64,
        min_chunk_size: int = 100
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size
        self.sentence_pattern = re.compile(r'(?<=[。！？.!?])\s*|\n\n+')
    
    def chunk(self, text: str, doc_id: str) -> List[Chunk]:
        if not text.strip():
            return []
        
        text = self._preprocess(text)
        sentences = self._split_sentences(text)
        
        if not sentences:
            return []
        
        return self._merge_into_chunks(sentences, doc_id)
    
    def _preprocess(self, text: str) -> str:
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()
    
    def _split_sentences(self, text: str) -> List[str]:
        paragraphs = text.split('\n\n')
        sentences = []
        for para in paragraphs:
            if not para.strip():
                continue
            parts = self.sentence_pattern.split(para)
            for part in parts:
                part = part.strip()
                if part:
                    sentences.append(part)
        return sentences
    
    def _merge_into_chunks(self, sentences: List[str], doc_id: str) -> List[Chunk]:
        chunks = []
        current_chunk_sentences = []
        current_size = 0
        current_start = 0
        position = 0
        
        for sentence in sentences:
            sentence_len = len(sentence)
            
            if current_size + sentence_len > self.chunk_size and current_chunk_sentences:
                chunk_content = " ".join(current_chunk_sentences)
                chunks.append(Chunk(
                    document_id=doc_id,
                    content=chunk_content,
                    start_pos=current_start,
                    end_pos=position,
                    metadata={"sentence_count": len(current_chunk_sentences)}
                ))
                
                overlap_sentences = []
                overlap_size = 0
                for s in reversed(current_chunk_sentences):
                    if overlap_size + len(s) <= self.chunk_overlap:
                        overlap_sentences.insert(0, s)
                        overlap_size += len(s)
                    else:
                        break
                
                current_chunk_sentences = overlap_sentences
                current_size = overlap_size
                current_start = position - overlap_size
            
            current_chunk_sentences.append(sentence)
            current_size += sentence_len
            position += sentence_len + 1
        
        if current_chunk_sentences:
            chunk_content = " ".join(current_chunk_sentences)
            if len(chunk_content) >= self.min_chunk_size or not chunks:
                chunks.append(Chunk(
                    document_id=doc_id,
                    content=chunk_content,
                    start_pos=current_start,
                    end_pos=position,
                    metadata={"sentence_count": len(current_chunk_sentences)}
                ))
            elif chunks:
                prev_chunk = chunks[-1]
                prev_chunk.content += " " + chunk_content
                prev_chunk.end_pos = position
                prev_chunk.metadata["sentence_count"] += len(current_chunk_sentences)
        
        return chunks


class RecursiveChunker(TextChunker):
    """
    递归分块器
    
    备选分块方案，使用递归策略处理长文本
    """
    
    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 64,
        separators: List[str] = None
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = separators or ["\n\n", "\n", "。", ".", " ", ""]
    
    def chunk(self, text: str, doc_id: str) -> List[Chunk]:
        raw_chunks = self._split_recursive(text, self.separators)
        
        chunks = []
        position = 0
        for i, content in enumerate(raw_chunks):
            chunks.append(Chunk(
                document_id=doc_id,
                content=content,
                start_pos=position,
                end_pos=position + len(content),
                metadata={"chunk_index": i}
            ))
            position += len(content)
        
        return chunks
    
    def _split_recursive(self, text: str, separators: List[str]) -> List[str]:
        if not text:
            return []
        
        if len(text) <= self.chunk_size:
            return [text]
        
        if not separators:
            return self._force_split(text)
        
        separator = separators[0]
        remaining_separators = separators[1:]
        
        if separator == "":
            return self._force_split(text)
        
        parts = text.split(separator)
        chunks = []
        current = ""
        
        for part in parts:
            test = current + separator + part if current else part
            
            if len(test) <= self.chunk_size:
                current = test
            else:
                if current:
                    chunks.append(current)
                
                if len(part) > self.chunk_size:
                    sub_chunks = self._split_recursive(part, remaining_separators)
                    chunks.extend(sub_chunks)
                    current = ""
                else:
                    current = part
        
        if current:
            chunks.append(current)
        
        return chunks
    
    def _force_split(self, text: str) -> List[str]:
        chunks = []
        for i in range(0, len(text), self.chunk_size - self.chunk_overlap):
            chunk = text[i:i + self.chunk_size]
            if chunk:
                chunks.append(chunk)
        return chunks


class DocumentProcessor:
    """
    文档处理器
    
    负责加载、解析和分块文档
    """
    
    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 64,
        supported_formats: List[str] = None
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.supported_formats = supported_formats or [".pdf", ".txt", ".docx", ".md"]
        self.loader = DocumentLoader(self.supported_formats)
        self.chunker = SemanticChunker(chunk_size, chunk_overlap)
    
    def load_document(self, file_path: str) -> Document:
        """加载文档"""
        return self.loader.load(file_path)
    
    def process(self, document: Document) -> List[Chunk]:
        """处理文档，执行分块"""
        chunks = self.chunker.chunk(document.content, document.id)
        
        for chunk in chunks:
            chunk.metadata.update({
                "filename": document.filename,
                "format": document.format,
                "document_created_at": document.created_at.isoformat(),
            })
        
        document.chunk_count = len(chunks)
        return chunks
    
    def process_file(self, file_path: str) -> Tuple[Document, List[Chunk]]:
        """处理文件的便捷方法"""
        document = self.load_document(file_path)
        chunks = self.process(document)
        return document, chunks
