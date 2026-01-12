"""
文档处理模块

负责文档解析、文本提取和智能分块
支持PDF、TXT、DOCX、Markdown格式
"""

import re
import uuid
import logging
from pathlib import Path
from dataclasses import dataclass, field
from datetime import datetime
from typing import List, Dict, Optional, Tuple

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """文档数据类"""
    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    filename: str = ""
    content: str = ""
    format: str = ""
    metadata: Dict = field(default_factory=dict)
    created_at: datetime = field(default_factory=datetime.now)
    chunk_count: int = 0


@dataclass
class Chunk:
    """文本块数据类"""
    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    document_id: str = ""
    content: str = ""
    start_pos: int = 0
    end_pos: int = 0
    metadata: Dict = field(default_factory=dict)


class DocumentParseError(Exception):
    """文档解析错误"""
    pass


class DocumentProcessor:
    """
    文档处理器
    
    负责加载、解析和分块文档
    """
    
    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 64,
        supported_formats: List[str] = None
    ):
        """
        初始化文档处理器
        
        Args:
            chunk_size: 分块大小（字符数）
            chunk_overlap: 分块重叠
            supported_formats: 支持的文件格式列表
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.supported_formats = supported_formats or [".pdf", ".txt", ".docx", ".md"]
        self.chunker = SemanticChunker(chunk_size, chunk_overlap)
    
    def load_document(self, file_path: str) -> Document:
        """
        加载并解析文档
        
        Args:
            file_path: 文档路径
            
        Returns:
            Document: 解析后的文档对象
            
        Raises:
            FileNotFoundError: 文件不存在
            DocumentParseError: 解析失败
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        suffix = path.suffix.lower()
        if suffix not in self.supported_formats:
            raise DocumentParseError(f"不支持的文件格式: {suffix}")
        
        # 根据格式选择解析器
        parsers = {
            ".pdf": self._parse_pdf,
            ".txt": self._parse_txt,
            ".docx": self._parse_docx,
            ".md": self._parse_markdown,
        }
        
        try:
            content = parsers[suffix](str(path))
        except Exception as e:
            raise DocumentParseError(f"解析文档失败: {e}")
        
        return Document(
            filename=path.name,
            content=content,
            format=suffix,
            metadata={
                "file_path": str(path),
                "file_size": path.stat().st_size,
            }
        )
    
    def _parse_pdf(self, file_path: str) -> str:
        """解析PDF文档"""
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n\n".join(text_parts)
        except ImportError:
            logger.warning("pypdf未安装，尝试使用备选方案")
            return self._parse_txt(file_path)
    
    def _parse_txt(self, file_path: str) -> str:
        """解析纯文本文档"""
        encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise DocumentParseError("无法识别文件编码")
    
    def _parse_docx(self, file_path: str) -> str:
        """解析Word文档"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            return "\n\n".join(paragraphs)
        except ImportError:
            logger.warning("python-docx未安装")
            raise DocumentParseError("需要安装python-docx来解析DOCX文件")
    
    def _parse_markdown(self, file_path: str) -> str:
        """解析Markdown文档"""
        content = self._parse_txt(file_path)
        # 简单的Markdown转纯文本处理
        # 移除代码块标记
        content = re.sub(r"```[\s\S]*?```", "", content)
        # 移除行内代码
        content = re.sub(r"`[^`]+`", "", content)
        # 移除图片链接
        content = re.sub(r"!\[.*?\]\(.*?\)", "", content)
        # 简化链接
        content = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", content)
        # 移除标题标记
        content = re.sub(r"^#+\s+", "", content, flags=re.MULTILINE)
        return content
    
    def process(self, document: Document) -> List[Chunk]:
        """
        处理文档，执行分块
        
        Args:
            document: 文档对象
            
        Returns:
            List[Chunk]: 文本块列表
        """
        chunks = self.chunker.chunk(document.content, document.id)
        
        # 添加文档元数据到每个块
        for chunk in chunks:
            chunk.metadata.update({
                "filename": document.filename,
                "format": document.format,
                "document_created_at": document.created_at.isoformat(),
            })
        
        document.chunk_count = len(chunks)
        return chunks
    
    def process_file(self, file_path: str) -> Tuple[Document, List[Chunk]]:
        """
        处理文件的便捷方法
        
        Args:
            file_path: 文件路径
            
        Returns:
            Tuple[Document, List[Chunk]]: 文档和文本块列表
        """
        document = self.load_document(file_path)
        chunks = self.process(document)
        return document, chunks


class SemanticChunker:
    """
    语义分块器 - 创新点
    
    基于语义边界进行智能分块，避免在句子中间切分
    """
    
    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 64,
        min_chunk_size: int = 100
    ):
        """
        初始化语义分块器
        
        Args:
            chunk_size: 目标块大小
            chunk_overlap: 块重叠大小
            min_chunk_size: 最小块大小
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size
        
        # 句子分割正则表达式
        # 支持中英文句子结束符
        self.sentence_pattern = re.compile(
            r'(?<=[。！？.!?])\s*|\n\n+'
        )
    
    def chunk(self, text: str, doc_id: str) -> List[Chunk]:
        """
        对文本进行语义分块
        
        算法：
        1. 先按段落分割
        2. 对长段落按句子分割
        3. 合并小片段直到达到目标大小
        4. 保持语义边界完整性
        
        Args:
            text: 要分块的文本
            doc_id: 文档ID
            
        Returns:
            List[Chunk]: 文本块列表
        """
        if not text.strip():
            return []
        
        # 预处理文本
        text = self._preprocess(text)
        
        # 分割成句子
        sentences = self._split_sentences(text)
        
        if not sentences:
            return []
        
        # 合并句子成块
        chunks = self._merge_into_chunks(sentences, doc_id)
        
        return chunks
    
    def _preprocess(self, text: str) -> str:
        """预处理文本"""
        # 规范化空白字符
        text = re.sub(r'\s+', ' ', text)
        # 规范化换行
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()
    
    def _split_sentences(self, text: str) -> List[str]:
        """
        将文本分割成句子
        
        Args:
            text: 输入文本
            
        Returns:
            List[str]: 句子列表
        """
        # 先按段落分割
        paragraphs = text.split('\n\n')
        
        sentences = []
        for para in paragraphs:
            if not para.strip():
                continue
            
            # 按句子分割
            parts = self.sentence_pattern.split(para)
            for part in parts:
                part = part.strip()
                if part:
                    sentences.append(part)
        
        return sentences
    
    def _merge_into_chunks(
        self, 
        sentences: List[str], 
        doc_id: str
    ) -> List[Chunk]:
        """
        将句子合并成文本块
        
        Args:
            sentences: 句子列表
            doc_id: 文档ID
            
        Returns:
            List[Chunk]: 文本块列表
        """
        chunks = []
        current_chunk_sentences = []
        current_size = 0
        current_start = 0
        position = 0
        
        for sentence in sentences:
            sentence_len = len(sentence)
            
            # 如果当前块加上新句子超过目标大小
            if current_size + sentence_len > self.chunk_size and current_chunk_sentences:
                # 创建当前块
                chunk_content = " ".join(current_chunk_sentences)
                chunks.append(Chunk(
                    document_id=doc_id,
                    content=chunk_content,
                    start_pos=current_start,
                    end_pos=position,
                    metadata={"sentence_count": len(current_chunk_sentences)}
                ))
                
                # 处理重叠：保留最后几个句子
                overlap_sentences = []
                overlap_size = 0
                for s in reversed(current_chunk_sentences):
                    if overlap_size + len(s) <= self.chunk_overlap:
                        overlap_sentences.insert(0, s)
                        overlap_size += len(s)
                    else:
                        break
                
                current_chunk_sentences = overlap_sentences
                current_size = overlap_size
                current_start = position - overlap_size
            
            current_chunk_sentences.append(sentence)
            current_size += sentence_len
            position += sentence_len + 1  # +1 for space
        
        # 处理最后一个块
        if current_chunk_sentences:
            chunk_content = " ".join(current_chunk_sentences)
            if len(chunk_content) >= self.min_chunk_size or not chunks:
                chunks.append(Chunk(
                    document_id=doc_id,
                    content=chunk_content,
                    start_pos=current_start,
                    end_pos=position,
                    metadata={"sentence_count": len(current_chunk_sentences)}
                ))
            elif chunks:
                # 太短则合并到前一个块
                prev_chunk = chunks[-1]
                prev_chunk.content += " " + chunk_content
                prev_chunk.end_pos = position
                prev_chunk.metadata["sentence_count"] += len(current_chunk_sentences)
        
        return chunks


class RecursiveChunker:
    """
    递归分块器
    
    备选分块方案，使用递归策略处理长文本
    """
    
    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 64,
        separators: List[str] = None
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = separators or ["\n\n", "\n", "。", ".", " ", ""]
    
    def chunk(self, text: str, doc_id: str) -> List[Chunk]:
        """递归分块"""
        raw_chunks = self._split_recursive(text, self.separators)
        
        chunks = []
        position = 0
        for i, content in enumerate(raw_chunks):
            chunks.append(Chunk(
                document_id=doc_id,
                content=content,
                start_pos=position,
                end_pos=position + len(content),
                metadata={"chunk_index": i}
            ))
            position += len(content)
        
        return chunks
    
    def _split_recursive(
        self, 
        text: str, 
        separators: List[str]
    ) -> List[str]:
        """递归分割文本"""
        if not text:
            return []
        
        if len(text) <= self.chunk_size:
            return [text]
        
        if not separators:
            # 没有分隔符了，强制分割
            return self._force_split(text)
        
        separator = separators[0]
        remaining_separators = separators[1:]
        
        if separator == "":
            return self._force_split(text)
        
        parts = text.split(separator)
        
        chunks = []
        current = ""
        
        for part in parts:
            test = current + separator + part if current else part
            
            if len(test) <= self.chunk_size:
                current = test
            else:
                if current:
                    chunks.append(current)
                
                if len(part) > self.chunk_size:
                    # 递归处理大块
                    sub_chunks = self._split_recursive(part, remaining_separators)
                    chunks.extend(sub_chunks)
                    current = ""
                else:
                    current = part
        
        if current:
            chunks.append(current)
        
        return chunks
    
    def _force_split(self, text: str) -> List[str]:
        """强制按大小分割"""
        chunks = []
        for i in range(0, len(text), self.chunk_size - self.chunk_overlap):
            chunk = text[i:i + self.chunk_size]
            if chunk:
                chunks.append(chunk)
        return chunks
